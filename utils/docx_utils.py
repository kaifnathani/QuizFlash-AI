from io import BytesIO
from typing import List
from docx import Document


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file bytes using python-docx."""
    try:
        doc = Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

        # also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n".join(paragraphs)
    except Exception:
        return ""
